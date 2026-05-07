"""Genera 2 contratos DOCX y los convierte a PDF.

Contrato 1: Arrendamiento Vivienda - Bogotá
  Arrendador: María Elena Rodríguez (cliente Premium del HOL)
  Arrendatario: Carlos Andrés Moreno (también cliente Estándar)
  
Contrato 2: Arrendamiento Comercial - Cali
  Arrendador: Ana Patricia Silva (cliente Premium del HOL)
  Arrendatario: Diana Carolina Pérez (cliente Estándar del HOL)
  
Ambos contratos están enlazados con la tabla CLIENTES y POLIZAS del Workshop.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE


def build_contract(out_path, data):
    doc = Document()
    # Margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    # Default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    # Title
    title = doc.add_heading(data['titulo'], level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(data['subtitulo']).italic = True
    
    doc.add_paragraph()
    
    # === PÁGINA 1: Identificación de las partes ===
    doc.add_heading('PRIMERA. IDENTIFICACIÓN DE LAS PARTES', level=1)
    p = doc.add_paragraph()
    p.add_run(f"Entre los suscritos a saber: ")
    p.add_run(data['arrendador_nombre']).bold = True
    p.add_run(f", mayor de edad, identificad{data['arrendador_genero']} con cédula de ciudadanía número ")
    p.add_run(data['arrendador_cedula']).bold = True
    p.add_run(f" expedida en {data['arrendador_ciudad_exp']}, con domicilio en {data['arrendador_domicilio']}, ")
    p.add_run(f"quien en adelante y para efectos del presente contrato se denominará ")
    p.add_run('EL ARRENDADOR').bold = True
    p.add_run(f", por una parte; y por la otra, ")
    p.add_run(data['arrendatario_nombre']).bold = True
    p.add_run(f", mayor de edad, identificad{data['arrendatario_genero']} con cédula de ciudadanía número ")
    p.add_run(data['arrendatario_cedula']).bold = True
    p.add_run(f" expedida en {data['arrendatario_ciudad_exp']}, con domicilio en {data['arrendatario_domicilio']}, ")
    p.add_run(f"quien en adelante se denominará ")
    p.add_run('EL ARRENDATARIO').bold = True
    p.add_run(f", hemos decidido celebrar el presente CONTRATO DE {data['tipo_arr'].upper()}, ")
    p.add_run("el cual se regirá por las cláusulas y condiciones que a continuación se detallan, en concordancia con la Ley 820 de 2003 y demás normas concordantes del Código Civil Colombiano.")
    
    doc.add_paragraph()
    
    doc.add_heading('SEGUNDA. OBJETO DEL CONTRATO', level=1)
    p = doc.add_paragraph(
        f"EL ARRENDADOR entrega a título de arrendamiento a EL ARRENDATARIO, y este recibe a satisfacción, "
        f"el inmueble identificado a continuación:"
    )
    
    doc.add_paragraph(f"• Tipo de inmueble: {data['inmueble_tipo']}", style='List Bullet')
    doc.add_paragraph(f"• Dirección: {data['inmueble_direccion']}", style='List Bullet')
    doc.add_paragraph(f"• Ciudad: {data['inmueble_ciudad']}", style='List Bullet')
    doc.add_paragraph(f"• Área aproximada: {data['inmueble_area']}", style='List Bullet')
    doc.add_paragraph(f"• Estrato socioeconómico: {data['inmueble_estrato']}", style='List Bullet')
    doc.add_paragraph(f"• Matrícula inmobiliaria: {data['inmueble_matricula']}", style='List Bullet')
    
    p = doc.add_paragraph()
    p.add_run(
        f"\nEL ARRENDATARIO destinará el inmueble exclusivamente para {data['destinacion']}, "
        f"y se obliga a no cambiar dicha destinación sin el consentimiento previo y por escrito de EL ARRENDADOR. "
        f"Cualquier cambio de uso constituirá causal de terminación inmediata del contrato."
    )
    
    doc.add_page_break()
    
    # === PÁGINA 2: Canon, plazo, depósito ===
    doc.add_heading('TERCERA. CANON DE ARRENDAMIENTO', level=1)
    p = doc.add_paragraph()
    p.add_run(f"El canon mensual de arrendamiento se establece en la suma de ")
    p.add_run(f"{data['canon_letras']} (${data['canon_numero']}) MONEDA CORRIENTE").bold = True
    p.add_run(f", que EL ARRENDATARIO se obliga a pagar a EL ARRENDADOR de manera anticipada dentro de los primeros cinco (5) días de cada mes, ")
    p.add_run(f"mediante consignación en la cuenta bancaria {data['cuenta_banco']} del banco {data['banco']}, a nombre de EL ARRENDADOR.")
    
    doc.add_paragraph(
        f"El canon mensual será reajustado anualmente en la misma proporción del incremento del Índice de Precios al Consumidor (IPC) "
        f"certificado por el DANE para el año calendario inmediatamente anterior, reajuste que operará automáticamente al cumplirse cada anualidad del contrato, "
        f"sin necesidad de comunicación o notificación previa."
    )
    
    doc.add_heading('CUARTA. PLAZO Y VIGENCIA', level=1)
    p = doc.add_paragraph()
    p.add_run(f"El presente contrato tendrá una duración de ")
    p.add_run(f"{data['plazo_letras']} ({data['plazo_meses']}) meses").bold = True
    p.add_run(f", contados a partir del día ")
    p.add_run(f"{data['fecha_inicio']}").bold = True
    p.add_run(f" y finalizando el día ")
    p.add_run(f"{data['fecha_fin']}").bold = True
    p.add_run(
        ". El contrato podrá prorrogarse de mutuo acuerdo entre las partes, mediante comunicación escrita con una antelación no inferior a treinta (30) días "
        "a la fecha de terminación. En caso de no manifestarse renovación expresa, el contrato terminará automáticamente sin necesidad de requerimiento judicial."
    )
    
    doc.add_heading('QUINTA. DEPÓSITO Y GARANTÍAS', level=1)
    doc.add_paragraph(
        f"EL ARRENDATARIO entrega en este acto a EL ARRENDADOR la suma de {data['deposito_letras']} (${data['deposito_numero']}) "
        f"a título de depósito en garantía, suma que será devuelta al ARRENDATARIO al finalizar el contrato, una vez se haya verificado el buen estado del inmueble "
        f"y se encuentren a paz y salvo los servicios públicos, administración y demás obligaciones."
    )
    doc.add_paragraph(
        f"Adicionalmente, EL ARRENDATARIO presenta como codeudor solidario a {data['codeudor']}, "
        f"identificado con cédula de ciudadanía número {data['codeudor_cedula']}, quien firma este contrato en señal de aceptación de las obligaciones aquí contraídas."
    )
    
    doc.add_page_break()
    
    # === PÁGINA 3: Obligaciones ===
    doc.add_heading('SEXTA. OBLIGACIONES DEL ARRENDATARIO', level=1)
    obligaciones_arr = [
        "Pagar puntualmente el canon mensual de arrendamiento dentro de los primeros cinco (5) días de cada mes calendario.",
        "Pagar oportunamente los servicios públicos domiciliarios (agua, luz, gas natural, internet, telefonía, aseo y alumbrado público) y la cuota de administración cuando aplique.",
        "Cuidar el inmueble como un buen padre de familia, manteniéndolo en perfecto estado de aseo, conservación y funcionamiento.",
        "No realizar modificaciones, mejoras, demoliciones ni adecuaciones estructurales sin autorización previa y por escrito de EL ARRENDADOR.",
        "No subarrendar total o parcialmente el inmueble, ni ceder los derechos derivados de este contrato a terceros.",
        "Permitir a EL ARRENDADOR o a su representante el ingreso al inmueble previa cita y con notificación de al menos 48 horas, para verificar el estado de conservación.",
        f"Contratar y mantener vigente durante todo el plazo del contrato un seguro de {data['seguro_obligatorio']} con cobertura mínima del 80% del valor comercial del inmueble.",
        "Restituir el inmueble a la terminación del contrato en las mismas condiciones en que lo recibió, salvo el deterioro natural por su uso adecuado.",
        "Reportar inmediatamente a EL ARRENDADOR cualquier daño, fuga, falla eléctrica o desperfecto que afecte la habitabilidad del inmueble.",
        "Cumplir con los reglamentos de propiedad horizontal y normas de convivencia del conjunto o edificio.",
    ]
    for o in obligaciones_arr:
        doc.add_paragraph(o, style='List Number')
    
    doc.add_page_break()
    
    # === PÁGINA 4: Prohibiciones, terminación, vehiculares ===
    doc.add_heading('SÉPTIMA. OBLIGACIONES DEL ARRENDADOR', level=1)
    obligaciones_dor = [
        "Entregar el inmueble en óptimas condiciones de habitabilidad y funcionamiento desde el primer día del contrato.",
        "Garantizar a EL ARRENDATARIO el goce pacífico y tranquilo del inmueble durante toda la vigencia del contrato.",
        "Realizar las reparaciones locativas y estructurales que no sean responsabilidad del ARRENDATARIO.",
        "Mantener al día el pago del impuesto predial y demás obligaciones tributarias del inmueble.",
        "Devolver el depósito en garantía al finalizar el contrato, dentro de los treinta (30) días siguientes a la entrega del inmueble.",
        f"Suscribir y mantener vigente la póliza de {data['poliza_arrendador']} con la aseguradora {data['aseguradora']}.",
    ]
    for o in obligaciones_dor:
        doc.add_paragraph(o, style='List Number')
    
    doc.add_heading('OCTAVA. CAUSALES DE TERMINACIÓN', level=1)
    doc.add_paragraph(
        "El presente contrato podrá darse por terminado de manera anticipada en los siguientes casos:"
    )
    causales = [
        "Mora en el pago de dos (2) o más cánones consecutivos de arrendamiento.",
        "Mora en el pago de servicios públicos por más de sesenta (60) días.",
        "Cambio de destinación del inmueble sin autorización del ARRENDADOR.",
        "Subarriendo o cesión no autorizada del contrato.",
        "Daños graves al inmueble por culpa o negligencia del ARRENDATARIO.",
        "Incumplimiento reiterado de las obligaciones aquí pactadas.",
        "Mutuo acuerdo entre las partes, formalizado por escrito.",
        "Las demás causales previstas en la Ley 820 de 2003.",
    ]
    for c in causales:
        doc.add_paragraph(c, style='List Bullet')
    
    doc.add_heading('NOVENA. CLÁUSULA PENAL', level=1)
    doc.add_paragraph(
        f"En caso de incumplimiento de cualquiera de las obligaciones contractuales por parte de EL ARRENDATARIO, "
        f"este pagará a EL ARRENDADOR a título de cláusula penal, sin perjuicio de los demás derechos legales, "
        f"una suma equivalente a {data['clausula_penal_meses']} cánones mensuales de arrendamiento, vigentes al momento del incumplimiento."
    )
    
    doc.add_page_break()
    
    # === PÁGINA 5: Domicilio, firmas ===
    doc.add_heading('DÉCIMA. DOMICILIO CONTRACTUAL Y NOTIFICACIONES', level=1)
    doc.add_paragraph(
        f"Para todos los efectos legales, las partes señalan como domicilio contractual la ciudad de {data['domicilio_contractual']}, Colombia. "
        f"Las notificaciones que deban hacerse en virtud del presente contrato se entregarán en las siguientes direcciones:"
    )
    p = doc.add_paragraph()
    p.add_run("EL ARRENDADOR: ").bold = True
    p.add_run(f"{data['arrendador_domicilio']} | Email: {data['arrendador_email']} | Teléfono: {data['arrendador_telefono']}")
    p = doc.add_paragraph()
    p.add_run("EL ARRENDATARIO: ").bold = True
    p.add_run(f"{data['arrendatario_domicilio']} | Email: {data['arrendatario_email']} | Teléfono: {data['arrendatario_telefono']}")
    
    doc.add_heading('DÉCIMA PRIMERA. RESOLUCIÓN DE CONTROVERSIAS', level=1)
    doc.add_paragraph(
        "Cualquier controversia que surja del presente contrato se resolverá mediante mecanismos alternativos de solución de conflictos. "
        "Las partes acuerdan acudir, en primera instancia, a un proceso de conciliación ante el Centro de Conciliación de la Cámara de Comercio competente. "
        "En caso de no llegar a un acuerdo, las controversias serán dirimidas por la jurisdicción ordinaria colombiana."
    )
    
    doc.add_heading('DÉCIMA SEGUNDA. PÓLIZA DE SEGUROS ASOCIADA', level=1)
    p = doc.add_paragraph()
    p.add_run(
        f"Las partes hacen constar que el inmueble objeto de este contrato cuenta con la siguiente póliza de seguros: "
    )
    p.add_run(f"\n• Tipo de póliza: {data['poliza_tipo']}").bold = False
    doc.add_paragraph(f"• Aseguradora: {data['aseguradora']}", style='List Bullet')
    doc.add_paragraph(f"• Número de póliza: {data['poliza_numero']}", style='List Bullet')
    doc.add_paragraph(f"• Cobertura total: ${data['cobertura']}", style='List Bullet')
    doc.add_paragraph(f"• Prima mensual: ${data['prima']}", style='List Bullet')
    doc.add_paragraph(f"• Vigencia: {data['poliza_vigencia']}", style='List Bullet')
    
    doc.add_heading('DÉCIMA TERCERA. ACEPTACIÓN Y FIRMAS', level=1)
    doc.add_paragraph(
        f"En constancia de lo anterior, las partes firman el presente contrato en dos (2) ejemplares del mismo tenor y valor, "
        f"en la ciudad de {data['domicilio_contractual']}, a los {data['fecha_firma']}."
    )
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run("________________________________________\n").bold = True
    p.add_run(f"{data['arrendador_nombre']}\n").bold = True
    p.add_run(f"C.C. {data['arrendador_cedula']}\n")
    p.add_run("EL ARRENDADOR").italic = True
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run("________________________________________\n").bold = True
    p.add_run(f"{data['arrendatario_nombre']}\n").bold = True
    p.add_run(f"C.C. {data['arrendatario_cedula']}\n")
    p.add_run("EL ARRENDATARIO").italic = True
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run("________________________________________\n").bold = True
    p.add_run(f"{data['codeudor']}\n").bold = True
    p.add_run(f"C.C. {data['codeudor_cedula']}\n")
    p.add_run("CODEUDOR SOLIDARIO").italic = True
    
    doc.save(out_path)
    print(f"✅ Generado: {out_path}")


# === CONTRATO 1: Vivienda Bogotá ===
contrato_1 = {
    'titulo': 'CONTRATO DE ARRENDAMIENTO DE VIVIENDA URBANA',
    'subtitulo': 'Regulado por la Ley 820 de 2003',
    
    'arrendador_nombre': 'MARÍA ELENA RODRÍGUEZ CASTILLO',
    'arrendador_genero': 'a',
    'arrendador_cedula': '52.874.369',
    'arrendador_ciudad_exp': 'Bogotá D.C.',
    'arrendador_domicilio': 'Calle 100 # 15-23, Apto 502, Bogotá D.C.',
    'arrendador_email': 'maria.rodriguez@correo.com',
    'arrendador_telefono': '+57 310 5847 932',
    
    'arrendatario_nombre': 'CARLOS ANDRÉS MORENO LÓPEZ',
    'arrendatario_genero': 'o',
    'arrendatario_cedula': '80.123.456',
    'arrendatario_ciudad_exp': 'Medellín',
    'arrendatario_domicilio': 'Carrera 15 # 93-47, Apto 304, Bogotá D.C.',
    'arrendatario_email': 'carlos.moreno@correo.com',
    'arrendatario_telefono': '+57 320 4521 786',
    
    'tipo_arr': 'arrendamiento de vivienda urbana',
    
    'inmueble_tipo': 'Apartamento residencial de tres (3) habitaciones, dos (2) baños, sala-comedor, cocina integral y balcón',
    'inmueble_direccion': 'Carrera 15 # 93-47, Apartamento 304, Edificio Mirador del Chicó',
    'inmueble_ciudad': 'Bogotá D.C., Colombia',
    'inmueble_area': '95 metros cuadrados (m²)',
    'inmueble_estrato': '5',
    'inmueble_matricula': '50C-1234567',
    'destinacion': 'vivienda familiar (uso exclusivamente residencial)',
    
    'canon_letras': 'DOS MILLONES OCHOCIENTOS MIL PESOS M/CTE',
    'canon_numero': '2.800.000',
    'cuenta_banco': '12345678901',
    'banco': 'Bancolombia',
    
    'plazo_letras': 'DOCE',
    'plazo_meses': '12',
    'fecha_inicio': '01 de noviembre de 2025',
    'fecha_fin': '31 de octubre de 2026',
    
    'deposito_letras': 'CINCO MILLONES SEISCIENTOS MIL PESOS M/CTE',
    'deposito_numero': '5.600.000',
    'codeudor': 'JORGE ALBERTO MORENO RUIZ',
    'codeudor_cedula': '79.456.321',
    
    'seguro_obligatorio': 'contenidos y responsabilidad civil',
    'poliza_arrendador': 'Protección Total Hogar',
    'aseguradora': 'Seguros Confianza S.A.',
    'clausula_penal_meses': 'tres (3)',
    
    'domicilio_contractual': 'Bogotá D.C.',
    'fecha_firma': '05 días del mes de noviembre del año dos mil veinticinco (2025)',
    
    'poliza_tipo': 'Hogar - Protección Total',
    'poliza_numero': 'POL-2025-HOG-002',
    'cobertura': '80.000.000',
    'prima': '120.000',
    'poliza_vigencia': 'Del 12 de noviembre de 2025 al 12 de noviembre de 2026',
}

# === CONTRATO 2: Local Comercial Cali ===
contrato_2 = {
    'titulo': 'CONTRATO DE ARRENDAMIENTO DE LOCAL COMERCIAL',
    'subtitulo': 'Regulado por el Código de Comercio Colombiano y la Ley 820 de 2003',
    
    'arrendador_nombre': 'ANA PATRICIA SILVA MENDOZA',
    'arrendador_genero': 'a',
    'arrendador_cedula': '31.789.234',
    'arrendador_ciudad_exp': 'Manizales',
    'arrendador_domicilio': 'Avenida 6N # 25-78, Casa 12, Cali, Valle del Cauca',
    'arrendador_email': 'ana.silva@correo.com',
    'arrendador_telefono': '+57 315 7894 256',
    
    'arrendatario_nombre': 'DIANA CAROLINA PÉREZ GUTIÉRREZ',
    'arrendatario_genero': 'a',
    'arrendatario_cedula': '55.432.198',
    'arrendatario_ciudad_exp': 'Cali',
    'arrendatario_domicilio': 'Calle 5 # 38-90, Edificio Centro Empresarial Versalles, Cali',
    'arrendatario_email': 'diana.perez@correo.com',
    'arrendatario_telefono': '+57 318 6234 547',
    
    'tipo_arr': 'arrendamiento de local comercial',
    
    'inmueble_tipo': 'Local comercial en primer piso, con vitrina a la calle, baño privado, depósito y un parqueadero asignado',
    'inmueble_direccion': 'Calle 5 # 38-90, Local 102, Centro Empresarial Versalles',
    'inmueble_ciudad': 'Cali, Valle del Cauca, Colombia',
    'inmueble_area': '60 metros cuadrados (m²)',
    'inmueble_estrato': '4',
    'inmueble_matricula': '370-9876543',
    'destinacion': 'la operación de un establecimiento comercial dedicado a la venta de productos de belleza, accesorios y servicios de estética',
    
    'canon_letras': 'TRES MILLONES DOSCIENTOS MIL PESOS M/CTE',
    'canon_numero': '3.200.000',
    'cuenta_banco': '98765432101',
    'banco': 'Davivienda',
    
    'plazo_letras': 'DIECIOCHO',
    'plazo_meses': '18',
    'fecha_inicio': '15 de noviembre de 2025',
    'fecha_fin': '14 de mayo de 2027',
    
    'deposito_letras': 'NUEVE MILLONES SEISCIENTOS MIL PESOS M/CTE',
    'deposito_numero': '9.600.000',
    'codeudor': 'LUIS ENRIQUE PÉREZ RAMÍREZ',
    'codeudor_cedula': '16.234.987',
    
    'seguro_obligatorio': 'incendio, hurto y responsabilidad civil extracontractual',
    'poliza_arrendador': 'Protección Comercial Integral',
    'aseguradora': 'Seguros Confianza S.A.',
    'clausula_penal_meses': 'cinco (5)',
    
    'domicilio_contractual': 'Santiago de Cali',
    'fecha_firma': '18 días del mes de noviembre del año dos mil veinticinco (2025)',
    
    'poliza_tipo': 'Hogar - Protección Básica (Comercial)',
    'poliza_numero': 'POL-2025-HOG-004',
    'cobertura': '30.000.000',
    'prima': '55.000',
    'poliza_vigencia': 'Del 25 de noviembre de 2025 al 25 de noviembre de 2026',
}

import os
os.makedirs('/Users/jparrado/HOL-repo/HOL/AI_SUMMIT/datasets/documentos', exist_ok=True)

build_contract('/Users/jparrado/HOL-repo/HOL/AI_SUMMIT/datasets/documentos/CONTRATO_ARRENDAMIENTO_01.docx', contrato_1)
build_contract('/Users/jparrado/HOL-repo/HOL/AI_SUMMIT/datasets/documentos/CONTRATO_ARRENDAMIENTO_02.docx', contrato_2)
